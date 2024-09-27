
from docx import Document

doc = Document()

list1 = [
    ["  合同编号：3023", "登记时间：2023/9/23 ", "磅单ID:6578"],
    ["  姓名：李军", "   ", "车号：1"],
    ["  地址：份子地新源2队", "车型：四轮车"],
]


table1 = doc.add_table(rows=3, cols=3)
table1.cell(2, 0).merge(table1.cell(2, 1))

for row in range(len(list1)):
    cells = table1.rows[row].cells
    for col in range(len(list1[row])):
        cells[col].text = str(list1[row][col])

# table2 = doc.add_table(rows=4, cols=5)
# table1.cell(3, 2).merge(table1.cell(3, 5))
#
# # 设置表格样式（可选）
# table2.style = 'Table Grid'
#
# # 填充表格数据
# data = [
#     ["Header1", "Header2", "Header3", "Header4"],
#     ["Row1 Col1", "Row1 Col2", "Row1 Col3", "Row1 Col4"],
#     ["Row2 Col1", "Row2 Col2", "Row2 Col3", "Row2 Col4"]
# ]
#
# # 将数据填入表格
# for row_idx, row_data in enumerate(data):
#     row = table.rows[row_idx]
#     for col_idx, cell_data in enumerate(row_data):
#         cell = row.cells[col_idx]
#         cell.text = cell_data

doc.save("test.docx")
