from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from docx.enum.text import WD_TAB_ALIGNMENT


# 设置单元格的边框
def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "FF0000", "space": "0"},
        bottom={"sz": 12, "color": "00FF00", "val": "single"},
        left={"sz": 24, "val": "dashed", "shadow": "true"},
        right={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


# 设置单元格的文字格式
def set_cell_font(cell, format_dict: dict):
    for font_format in format_dict.keys():
        if font_format == '行间距':
            cell.paragraphs[0].paragraph_format.line_spacing = Pt(format_dict[font_format])  # 行间距，固定值0磅
        elif font_format == '首行缩进':
            cell.paragraphs[0].paragraph_format.first_line_indent = Pt(format_dict[font_format])  # 首行缩进0磅
        elif font_format == '段前间距':
            cell.paragraphs[0].paragraph_format.space_before = Pt(format_dict[font_format])  # 段前0磅
        elif font_format == '段后间距':
            cell.paragraphs[0].paragraph_format.space_after = Pt(format_dict[font_format])  # 段后0磅
        elif font_format == '文字居中':
            cell.paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
        elif font_format == '西文字体':
            cell.paragraphs[0].runs[0].font.name = format_dict[font_format]  # 西文字体
        elif font_format == '中文字体':
            cell.paragraphs[0].runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), format_dict[font_format])  # 中文字体
        elif font_format == '字体大小':
            cell.paragraphs[0].runs[0].font.size = Pt(format_dict[font_format])  # 字体大小
        elif font_format == '加粗':
            cell.paragraphs[0].runs[0].font.bold = format_dict[font_format]  # 是否加粗
        elif font_format == '斜体':
            cell.paragraphs[0].runs[0].font.italic = format_dict[font_format]  # 是否加粗
        elif font_format == '下划线':
            cell.paragraphs[0].runs[0].font.underline = format_dict[font_format]  # 是否下划线

