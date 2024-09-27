
from docx import Document
import pandas as pd
from utils.num_transfer import num_2_cn
from utils.word_table_utils import *
from docx.shared import Cm, Pt


def tickMakeServer(doc, row, comp_name):

    contract_id = row['合同号']
    reg_date = row['回皮时间'].split(' ')[0]
    weight_note_id = row['流水号']
    farmer_name = row['农户姓名']
    car_id = row['车牌号']
    address = row['地址']
    car_type = row['车辆类型']
    gross_weight = row['毛重']
    net_weight = row['净重']
    unit_price = row['单价']
    amount = row['金额']
    amount_cn = num_2_cn(str(row['金额']))
    registrant = row['登记人']  # 登记人  吴美娴
    weight_operator = row['重磅员']  # 司磅员

    table1 = doc.add_table(rows=4, cols=3)
    table1.cell(0, 0).merge(table1.cell(0, 1)).merge(table1.cell(0, 2))
    table1.cell(3, 0).merge(table1.cell(3, 1))

    # table1.cell(0, 0).text = "内蒙古伍星食品有限责任公司-收购磅单"
    table1.cell(0, 0).text = "{}-收购磅单".format(comp_name)
    set_cell_font(table1.cell(0, 0), {"文字居中": True, "加粗": True, "下划线": True})

    table1.cell(1, 0).text = "合同编号：{}".format(contract_id)
    table1.cell(1, 1).text = "回皮时间：{}".format(reg_date)
    table1.cell(1, 2).text = "磅单ID：{}".format(weight_note_id)
    table1.cell(2, 0).text = "姓名：{}".format(farmer_name)
    table1.cell(2, 1).text = "   "
    table1.cell(2, 2).text = "车号：{}".format(car_id)
    table1.cell(3, 0).text = "地址：{}".format(address)
    table1.cell(3, 2).text = "车型：{}".format(car_type)

    for tbl_row in table1.rows:
        for tbl_cell in tbl_row.cells:
            set_cell_font(tbl_cell, {"字体大小": 15, "段前间距": 0, "段后间距": 0})

    # 第二部分
    table2 = doc.add_table(rows=4, cols=5)
    table2.cell(0, 0).text = "产品名称"
    table2.cell(1, 0).text = "番茄"
    table2.cell(0, 1).text = "毛重（KG）"
    table2.cell(1, 1).text = str(gross_weight)
    table2.cell(0, 2).text = "净重（KG）"
    table2.cell(1, 2).text = str(net_weight)
    table2.cell(0, 3).text = "单价"
    table2.cell(1, 3).text = str(unit_price)
    table2.cell(0, 4).text = "金额"
    table2.cell(1, 4).text = str(amount)

    table2.cell(2, 1).merge(table2.cell(2, 2)).merge(table2.cell(2, 3)).merge(table2.cell(2, 4))
    table2.cell(2, 0).text = "金额大写："
    table2.cell(2, 1).text = amount_cn

    table2.cell(3, 0).merge(table2.cell(3, 1)).merge(table2.cell(3, 2)).merge(table2.cell(3, 3)).merge(table2.cell(3, 4))
    table2.cell(3, 0).text = "备注：盖章有效，遗失不补"
    table2.cell(3, 0).paragraphs[0].alignment = WD_TAB_ALIGNMENT.CENTER
    for tbl_row in table2.rows:
        for tbl_cell in tbl_row.cells:
            set_cell_font(tbl_cell, {"字体大小": 15, "段前间距": 0, "段后间距": 0})
            set_cell_border(tbl_cell,
                            top={"sz": 6, "color": "#000000", "val": "single"},
                            bottom={"sz": 6, "color": "#000000", "val": "single"},
                            left={"sz": 6, "color": "#000000", "val": "single"},
                            right={"sz": 6, "color": "#000000", "val": "single"}
                            )

    # 第三部分
    table3 = doc.add_table(rows=1, cols=3)
    table3.cell(0, 0).text = "登记人：{}".format(registrant)
    set_cell_font(table3.cell(0, 0), {"字体大小": 15})
    table3.cell(0, 1).text = "审核人："
    set_cell_font(table3.cell(0, 1), {"字体大小": 15})
    table3.cell(0, 2).text = "司磅员：{}".format(weight_operator)
    set_cell_font(table3.cell(0, 2), {"字体大小": 15})


if __name__ == '__main__':

    data = pd.read_excel("D:/私人/数据处理/data_process/数据处理结果.xlsx", dtype=str)
    data = data.head(8)
    data['登记人'] = '吴美娴'

    doc = Document()
    default_section = doc.sections[0]
    # 可直接修改宽度和高度，即纸张大小改为自定义（A4纸是21 * 29.7）
    default_section.page_width = Cm(21)
    default_section.page_height = Cm(29.7)

    # 修改页边距
    default_section.top_margin = Pt(26)  # 1CM = 28.35PT
    default_section.right_margin = Cm(1.2)
    default_section.bottom_margin = Pt(26)
    default_section.left_margin = Cm(1.2)

    for idx, row in data.iterrows():
        tickMakeServer(doc, row)
        print("完成第{}张票据生成".format(idx + 1))
        if (idx + 1) % 3 != 0:
            gap = doc.add_paragraph('------------------------------------------------------------------------------------------------------------------------------------------------')
            gap.paragraph_format.space_before = Pt(26)  # 段前10磅
            gap.paragraph_format.space_after = Pt(26)  # 段后10磅
        else:
            doc.add_page_break()

    doc.save("test.docx")

